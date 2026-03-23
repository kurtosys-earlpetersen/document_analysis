"""Word (DOCX) extraction using python-docx."""

from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import Block, ExtractedDocument


def _paragraph_to_block(p: Paragraph) -> Block | None:
    text = (p.text or "").strip()
    if not text:
        return None
    style = (p.style and p.style.name or "").lower()
    if "heading" in style or "title" in style:
        level = 1
        if "heading" in style:
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
        return Block(kind="heading", level=min(level, 6), text=text)
    return Block(kind="paragraph", text=text)


def _table_to_block(t: Table) -> Block:
    rows = [[c.text.strip() for c in row.cells] for row in t.rows]
    text = "\n".join(" | ".join(cell for cell in row) for row in rows)
    return Block(kind="table", text=text, metadata={"rows": len(rows), "cols": len(rows[0]) if rows else 0})


def parse_docx(path: Path) -> ExtractedDocument:
    doc = Document(path)
    blocks: list[Block] = []
    raw_chunks: list[str] = []

    for p in doc.paragraphs:
        b = _paragraph_to_block(p)
        if b:
            blocks.append(b)
            raw_chunks.append(b.text)
    for t in doc.tables:
        b = _table_to_block(t)
        blocks.append(b)
        raw_chunks.append(b.text)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            blocks.append(Block(kind="image", alt_text=None, metadata={"rel": rel.target_ref}))

    raw_text = "\n".join(raw_chunks)
    return ExtractedDocument(
        path=path,
        format="docx",
        raw_text=raw_text,
        blocks=blocks,
    )
